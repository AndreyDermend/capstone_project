"""
Automated test harness for the LexiAgent extraction pipeline.
Supports NDA, ConsultingAgreement, and EmploymentAgreement test prompts, scores extraction accuracy.

Scoring philosophy:
- Conservative extraction is fine — we don't penalize missing fields heavily
- Follow-up coverage is critical — every missing required field MUST get a question
- Follow-up question quality matters — questions must be clear and answerable
- Hallucination is the worst failure — extracting unsupported values is penalized hard

Usage:
    python app/test_extraction.py                    # run all NDA tests
    python app/test_extraction.py consulting         # run all consulting tests
    python app/test_extraction.py employment         # run all employment tests
    python app/test_extraction.py employment 1       # run employment test 1
    python app/test_extraction.py nda 1              # run NDA test 1
    python app/test_extraction.py all                # run all tests (all types)
"""

import json
import sys
import time
from pathlib import Path
from typing import Any, Dict, List, Optional

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "app"))

from run_intake_loop import (
    extract_answers_from_prompt,
    verify_and_prepare,
    add_derived_defaults,
    MODEL,
    required_fields,
)
from assemble_contract import assemble_contract, load_resources
from clause_rag import select_clauses_rag
from contract_docx import generate_contract_docx
from api_server import parse_follow_up_answers
from docx import Document as DocxDocument

OUTPUT_DIR = ROOT / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

ALL_CONTRACT_TYPES = [
    "NDA",
    "ConsultingAgreement",
    "EmploymentAgreement",
    "ServiceAgreement",
]

PARTY_FIELD_MAP = {
    "NDA": ["party_a_name", "party_b_name"],
    "ConsultingAgreement": ["client_name", "consultant_name"],
    "EmploymentAgreement": ["employer_name", "employee_name"],
    "ServiceAgreement": ["client_name", "contractor_name"],
}

# ---------------------------------------------------------------------------
# Test definitions from the handoff document (Section 16-17)
# ---------------------------------------------------------------------------
TESTS = [
    {
        "id": 1,
        "name": "Fully specified unilateral",
        "prompt": (
            "We need a unilateral NDA between Acme Corp., a Delaware corporation, "
            "and Beta Ventures LLC, a Connecticut LLC, to evaluate a potential partnership "
            "for 3 years under Connecticut law using litigation. Acme's email is "
            "legal@acme.com and Beta's email is ops@betaventures.com."
        ),
        "expected_extracted": {
            "nda_type": "Unilateral",
            "party_a_name": "Acme Corp.",
            "party_a_entity_details": "Delaware corporation",
            "party_a_email": "legal@acme.com",
            "party_b_name": "Beta Ventures LLC",
            "party_b_entity_details": "Connecticut LLC",
            "party_b_email": "ops@betaventures.com",
            "purpose": "evaluate a potential partnership",
            "confidentiality_period_number": 3,
            "confidentiality_period_unit": "years",
            "governing_law": "Connecticut",
            "dispute_resolution_method": "Litigation",
        },
        "expected_follow_ups": [],
        # Fields the user DID NOT mention — extracting these is hallucination
        "must_not_extract": [],
    },
    {
        "id": 2,
        "name": "Partially specified unilateral",
        "prompt": (
            "We need a unilateral NDA between Acme Corp and Beta Ventures LLC to evaluate "
            "a potential partnership for 3 years under Connecticut law using litigation."
        ),
        "expected_extracted": {
            "nda_type": "Unilateral",
            "party_a_name": "Acme Corp",
            "party_b_name": "Beta Ventures LLC",
            "purpose": "evaluate a potential partnership",
            "confidentiality_period_number": 3,
            "confidentiality_period_unit": "years",
            "governing_law": "Connecticut",
            "dispute_resolution_method": "Litigation",
        },
        "expected_follow_ups": [
            "party_a_entity_details",
            "party_a_email",
            "party_b_entity_details",
            "party_b_email",
        ],
        "must_not_extract": [
            "party_a_entity_details",
            "party_a_email",
            "party_b_entity_details",
            "party_b_email",
        ],
    },
    {
        "id": 3,
        "name": "Role-based unilateral (implies unilateral)",
        "prompt": (
            "Acme Corp wants to disclose confidential information to Beta Ventures LLC "
            "so they can evaluate a potential partnership. The arrangement should last 3 years, "
            "use Connecticut law, and resolve disputes through litigation. Acme's email is "
            "legal@acme.com and Beta's email is ops@betaventures.com."
        ),
        "expected_extracted": {
            "party_a_name": "Acme Corp",
            "party_a_email": "legal@acme.com",
            "party_b_name": "Beta Ventures LLC",
            "party_b_email": "ops@betaventures.com",
            "purpose": "evaluate a potential partnership",
            "confidentiality_period_number": 3,
            "confidentiality_period_unit": "years",
            "governing_law": "Connecticut",
            "dispute_resolution_method": "Litigation",
        },
        "expected_follow_ups": [
            "party_a_entity_details",
            "party_b_entity_details",
        ],
        "acceptable_extra_extracted": ["nda_type"],
        "must_not_extract": [
            "party_a_entity_details",
            "party_b_entity_details",
        ],
    },
    {
        "id": 4,
        "name": "Fully specified mutual",
        "prompt": (
            "We need a mutual NDA between Acme Corp., a Delaware corporation, and Beta Ventures LLC, "
            "a Connecticut LLC, to evaluate a potential partnership for 3 years under Connecticut law "
            "using litigation. Acme's email is legal@acme.com and Beta's email is ops@betaventures.com."
        ),
        "expected_extracted": {
            "nda_type": "Mutual",
            "party_a_name": "Acme Corp.",
            "party_a_entity_details": "Delaware corporation",
            "party_a_email": "legal@acme.com",
            "party_b_name": "Beta Ventures LLC",
            "party_b_entity_details": "Connecticut LLC",
            "party_b_email": "ops@betaventures.com",
            "purpose": "evaluate a potential partnership",
            "confidentiality_period_number": 3,
            "confidentiality_period_unit": "years",
            "governing_law": "Connecticut",
            "dispute_resolution_method": "Litigation",
        },
        "expected_follow_ups": [],
        "must_not_extract": [],
    },
    {
        "id": 5,
        "name": "Partially specified mutual",
        "prompt": (
            "We need a mutual NDA between Acme Corp and Beta Ventures LLC for evaluating a "
            "potential partnership for 3 years under Connecticut law."
        ),
        "expected_extracted": {
            "nda_type": "Mutual",
            "party_a_name": "Acme Corp",
            "party_b_name": "Beta Ventures LLC",
            "purpose": "evaluating a potential partnership",
            "confidentiality_period_number": 3,
            "confidentiality_period_unit": "years",
            "governing_law": "Connecticut",
        },
        "expected_follow_ups": [
            "party_a_entity_details",
            "party_a_email",
            "party_b_entity_details",
            "party_b_email",
            "dispute_resolution_method",
        ],
        "must_not_extract": [
            "party_a_entity_details",
            "party_a_email",
            "party_b_entity_details",
            "party_b_email",
            "dispute_resolution_method",
        ],
    },
    {
        "id": 6,
        "name": "Conversational / indirect",
        "prompt": (
            "Acme Corp and Beta Ventures LLC are exploring a possible partnership and need an NDA. "
            "It should last 3 years, follow Connecticut law, and use litigation if there's a dispute. "
            "Acme's email is legal@acme.com and Beta's email is ops@betaventures.com."
        ),
        "expected_extracted": {
            "party_a_name": "Acme Corp",
            "party_a_email": "legal@acme.com",
            "party_b_name": "Beta Ventures LLC",
            "party_b_email": "ops@betaventures.com",
            "purpose": "exploring a possible partnership",
            "confidentiality_period_number": 3,
            "confidentiality_period_unit": "years",
            "governing_law": "Connecticut",
            "dispute_resolution_method": "Litigation",
        },
        "expected_follow_ups": [
            "nda_type",
            "party_a_entity_details",
            "party_b_entity_details",
        ],
        "must_not_extract": [
            "nda_type",
            "party_a_entity_details",
            "party_b_entity_details",
        ],
    },
    # ------------------------------------------------------------------
    # Edge-case tests 7-10
    # ------------------------------------------------------------------
    {
        "id": 7,
        "name": "Typos and abbreviations",
        "prompt": (
            "We need a unilateral NDA btwn Acme Corp and Beta Ventures for evaluating "
            "a partnership, 3 yrs, Connecticut law, litigation. "
            "Acme email: legal@acme.com, Beta email: ops@betaventures.com."
        ),
        "expected_extracted": {
            "party_a_name": "Acme Corp",
            "party_b_name": "Beta Ventures",
            "purpose": "evaluating a partnership",
            "confidentiality_period_number": 3,
            "confidentiality_period_unit": "years",
            "governing_law": "Connecticut",
            "dispute_resolution_method": "Litigation",
            "party_a_email": "legal@acme.com",
            "party_b_email": "ops@betaventures.com",
        },
        "expected_follow_ups": [
            "party_a_entity_details",
            "party_b_entity_details",
        ],
        "acceptable_extra_extracted": ["nda_type"],
        "must_not_extract": [
            "party_a_entity_details",
            "party_b_entity_details",
        ],
    },
    {
        "id": 8,
        "name": "Extra irrelevant information",
        "prompt": (
            "We need a mutual NDA between Acme Corp., a Delaware corporation, and Beta Ventures LLC, "
            "a Connecticut LLC. Purpose is evaluating a potential partnership for 3 years under "
            "Connecticut law using arbitration. Acme's email is legal@acme.com and Beta's email is "
            "ops@betaventures.com. The CEO's name is John Smith, they prefer blue ink signatures, "
            "and the office is at 123 Main St."
        ),
        "expected_extracted": {
            "nda_type": "Mutual",
            "party_a_name": "Acme Corp.",
            "party_a_entity_details": "Delaware corporation",
            "party_a_email": "legal@acme.com",
            "party_b_name": "Beta Ventures LLC",
            "party_b_entity_details": "Connecticut LLC",
            "party_b_email": "ops@betaventures.com",
            "purpose": "evaluating a potential partnership",
            "confidentiality_period_number": 3,
            "confidentiality_period_unit": "years",
            "governing_law": "Connecticut",
            "dispute_resolution_method": "Arbitration",
        },
        "expected_follow_ups": [],
        "must_not_extract": [],
    },
    {
        "id": 9,
        "name": "Minimal prompt",
        "prompt": "NDA between Acme and Beta.",
        "expected_extracted": {
            "party_a_name": "Acme",
            "party_b_name": "Beta",
        },
        "expected_follow_ups": [
            "nda_type",
            "party_a_entity_details",
            "party_a_email",
            "party_b_entity_details",
            "party_b_email",
            "purpose",
            "confidentiality_period_number",
            "confidentiality_period_unit",
            "governing_law",
            "dispute_resolution_method",
        ],
        "must_not_extract": [
            "nda_type",
            "party_a_entity_details",
            "party_a_email",
            "party_b_entity_details",
            "party_b_email",
            "purpose",
            "confidentiality_period_number",
            "confidentiality_period_unit",
            "governing_law",
            "dispute_resolution_method",
        ],
    },
    {
        "id": 10,
        "name": "Conflicting info (mutual + disclose)",
        "prompt": (
            "We need a mutual NDA. Acme Corp wants to disclose confidential information to "
            "Beta Ventures LLC for evaluating a potential partnership. 3 years, Connecticut law, "
            "litigation. Acme's email is legal@acme.com and Beta's email is ops@betaventures.com."
        ),
        "expected_extracted": {
            "party_a_name": "Acme Corp",
            "party_b_name": "Beta Ventures LLC",
            "purpose": "evaluating a potential partnership",
            "confidentiality_period_number": 3,
            "confidentiality_period_unit": "years",
            "governing_law": "Connecticut",
            "dispute_resolution_method": "Litigation",
            "party_a_email": "legal@acme.com",
            "party_b_email": "ops@betaventures.com",
        },
        # nda_type is ambiguous — user says "mutual" but describes unilateral behavior
        # Conservative: should either extract "Mutual" (trusting explicit statement) or ask
        "expected_follow_ups": [
            "party_a_entity_details",
            "party_b_entity_details",
        ],
        "acceptable_extra_extracted": ["nda_type"],
        "must_not_extract": [
            "party_a_entity_details",
            "party_b_entity_details",
        ],
    },
]


# ---------------------------------------------------------------------------
# Consulting Agreement tests
# ---------------------------------------------------------------------------
CONSULTING_TESTS = [
    {
        "id": 1,
        "name": "Fully specified consulting agreement",
        "contract_type": "ConsultingAgreement",
        "prompt": (
            "We need a consulting agreement between Acme Corp at 123 Main St, Hartford CT 06103 "
            "(legal@acme.com) and Jane Smith Consulting at 456 Oak Ave, New Haven CT 06510 "
            "(jane@smithconsulting.com). Jane will provide marketing strategy and brand positioning "
            "services for $25,000, paid net 30. The agreement starts January 15, 2026 and ends "
            "June 30, 2026. Either party can terminate with 30 days notice. Connecticut law, "
            "arbitration for disputes. Client owns all IP."
        ),
        "expected_extracted": {
            "consulting_type": "Standard",
            "client_name": "Acme Corp",
            "client_address": "123 Main St, Hartford CT 06103",
            "client_email": "legal@acme.com",
            "consultant_name": "Jane Smith Consulting",
            "consultant_address": "456 Oak Ave, New Haven CT 06510",
            "consultant_email": "jane@smithconsulting.com",
            "effective_date": "January 15, 2026",
            "services_description": "marketing strategy and brand positioning",
            "compensation_amount": "$25,000",
            "payment_schedule": "Net 30",
            "term_end_date": "June 30, 2026",
            "termination_notice_days": 30,
            "governing_law": "Connecticut",
            "ip_ownership": "Client",
            "dispute_resolution_method": "Arbitration",
        },
        "expected_follow_ups": [],
        "must_not_extract": [],
    },
    {
        "id": 2,
        "name": "Partially specified consulting (missing addresses/emails/dates)",
        "contract_type": "ConsultingAgreement",
        "prompt": (
            "I need a consulting agreement. BrightWave Marketing will provide social media management "
            "and content creation for TechStart Inc. The fee is $5,000 per month. Use New York law. "
            "The consultant keeps the IP."
        ),
        "expected_extracted": {
            "consulting_type": "Standard",
            "client_name": "TechStart Inc",
            "consultant_name": "BrightWave Marketing",
            "services_description": "social media management and content creation",
            "compensation_amount": "$5,000",
            "governing_law": "New York",
            "ip_ownership": "Consultant",
        },
        "expected_follow_ups": [
            "client_address",
            "client_email",
            "consultant_address",
            "consultant_email",
            "effective_date",
            "term_end_date",
            "termination_notice_days",
        ],
        "acceptable_extra_extracted": ["payment_schedule"],
        "must_not_extract": [
            "client_address",
            "client_email",
            "consultant_address",
            "consultant_email",
            "effective_date",
            "term_end_date",
            "termination_notice_days",
        ],
    },
    {
        "id": 3,
        "name": "Conversational consulting prompt",
        "contract_type": "ConsultingAgreement",
        "prompt": (
            "We're hiring a consultant to help with our website redesign. The company is called "
            "PixelPerfect Design and they'll charge us $150 an hour. We want to use California law. "
            "My company is GlobalTech Solutions and my email is cto@globaltech.com."
        ),
        "expected_extracted": {
            "client_name": "GlobalTech Solutions",
            "client_email": "cto@globaltech.com",
            "consultant_name": "PixelPerfect Design",
            "services_description": "website redesign",
            "compensation_amount": "$150",
            "governing_law": "California",
        },
        "expected_follow_ups": [
            "client_address",
            "consultant_address",
            "consultant_email",
            "effective_date",
            "payment_schedule",
            "term_end_date",
            "termination_notice_days",
            "ip_ownership",
        ],
        "acceptable_extra_extracted": ["consulting_type"],
        "must_not_extract": [
            "client_address",
            "consultant_address",
            "consultant_email",
            "effective_date",
            "term_end_date",
            "termination_notice_days",
        ],
    },
    {
        "id": 4,
        "name": "Minimal consulting prompt",
        "contract_type": "ConsultingAgreement",
        "prompt": "Consulting agreement between DataPro Analytics and Summit Corp.",
        "expected_extracted": {
            "consultant_name": "DataPro Analytics",
            "client_name": "Summit Corp",
        },
        "expected_follow_ups": [
            "client_address",
            "client_email",
            "consultant_address",
            "consultant_email",
            "effective_date",
            "services_description",
            "compensation_amount",
            "payment_schedule",
            "term_end_date",
            "termination_notice_days",
            "governing_law",
            "ip_ownership",
        ],
        "acceptable_extra_extracted": ["consulting_type"],
        "must_not_extract": [
            "client_address",
            "client_email",
            "consultant_address",
            "consultant_email",
            "effective_date",
            "services_description",
            "compensation_amount",
            "payment_schedule",
            "term_end_date",
            "termination_notice_days",
            "governing_law",
            "ip_ownership",
        ],
    },
]


# ---------------------------------------------------------------------------
# Employment Agreement tests
# ---------------------------------------------------------------------------
EMPLOYMENT_TESTS = [
    {
        "id": 1,
        "name": "Fully specified employment agreement",
        "contract_type": "EmploymentAgreement",
        "prompt": (
            "We need an employment agreement for TechCorp Inc. at 100 Innovation Dr, San Francisco, CA 94105 "
            "hiring Jane Smith of 456 Oak Ave, San Francisco, CA 94110 as a Senior Software Engineer. "
            "She will design, develop, and maintain software applications. Full-time, starting January 15, 2026. "
            "Salary is $120,000 per year paid bi-weekly. She'll work at 100 Innovation Dr, San Francisco, CA 94105. "
            "30 days notice to terminate. California law, arbitration for disputes."
        ),
        "expected_extracted": {
            "employment_type": "Standard",
            "employer_name": "TechCorp Inc.",
            "employer_address": "100 Innovation Dr, San Francisco, CA 94105",
            "employee_name": "Jane Smith",
            "employee_address": "456 Oak Ave, San Francisco, CA 94110",
            "job_title": "Senior Software Engineer",
            "job_duties": "design, develop, and maintain software applications",
            "start_date": "January 15, 2026",
            "employment_basis": "Full-Time",
            "compensation_amount": "$120,000 per year",
            "pay_frequency": "Bi-Weekly",
            "work_location": "100 Innovation Dr, San Francisco, CA 94105",
            "termination_notice_days": 30,
            "governing_law": "California",
            "dispute_resolution_method": "Arbitration",
        },
        "expected_follow_ups": [],
        "must_not_extract": [],
    },
    {
        "id": 2,
        "name": "Partially specified employment (missing addresses/duties)",
        "contract_type": "EmploymentAgreement",
        "prompt": (
            "I need to hire a Marketing Manager at Bright Solutions LLC. The salary is "
            "$85,000 per year, paid semi-monthly. The job starts on March 1, 2026. "
            "We're in Texas. 14 days notice to terminate."
        ),
        "expected_extracted": {
            "employer_name": "Bright Solutions LLC",
            "job_title": "Marketing Manager",
            "start_date": "March 1, 2026",
            "compensation_amount": "$85,000",
            "pay_frequency": "Semi-Monthly",
            "governing_law": "Texas",
            "termination_notice_days": 14,
        },
        "expected_follow_ups": [
            "employer_address",
            "employee_name",
            "employee_address",
            "job_duties",
            "employment_basis",
            "work_location",
        ],
        "acceptable_extra_extracted": ["employment_type"],
        "must_not_extract": [
            "employer_address",
            "employee_name",
            "employee_address",
            "job_duties",
            "work_location",
        ],
    },
    {
        "id": 3,
        "name": "Conversational employment prompt",
        "contract_type": "EmploymentAgreement",
        "prompt": (
            "We're bringing on a new data analyst at our company, DataFlow Analytics. "
            "They'll be working part-time, $35 an hour, paid weekly. We need this under "
            "New York law. 30 days notice for termination."
        ),
        "expected_extracted": {
            "employer_name": "DataFlow Analytics",
            "job_title": "Data Analyst",
            "employment_basis": "Part-Time",
            "compensation_amount": "$35",
            "pay_frequency": "Weekly",
            "governing_law": "New York",
            "termination_notice_days": 30,
        },
        "expected_follow_ups": [
            "employer_address",
            "employee_name",
            "employee_address",
            "job_duties",
            "start_date",
            "work_location",
        ],
        "acceptable_extra_extracted": ["employment_type"],
        "must_not_extract": [
            "employer_address",
            "employee_name",
            "employee_address",
            "job_duties",
            "start_date",
            "work_location",
        ],
    },
    {
        "id": 4,
        "name": "Minimal employment prompt",
        "contract_type": "EmploymentAgreement",
        "prompt": "Employment agreement between Acme Inc and John Doe.",
        "expected_extracted": {
            "employer_name": "Acme Inc",
            "employee_name": "John Doe",
        },
        "expected_follow_ups": [
            "employer_address",
            "employee_address",
            "job_title",
            "job_duties",
            "start_date",
            "employment_basis",
            "compensation_amount",
            "pay_frequency",
            "work_location",
            "termination_notice_days",
            "governing_law",
        ],
        "acceptable_extra_extracted": ["employment_type"],
        "must_not_extract": [
            "employer_address",
            "employee_address",
            "job_title",
            "job_duties",
            "start_date",
            "compensation_amount",
            "pay_frequency",
            "work_location",
            "termination_notice_days",
            "governing_law",
        ],
    },
]

SERVICE_TESTS = [
    {
        "id": 1,
        "name": "Fully specified service agreement",
        "contract_type": "ServiceAgreement",
        "prompt": (
            "We need a service agreement between GlobalTech Solutions at 500 Market St, San Francisco CA 94105 "
            "and CleanPro Services LLC at 200 Oak Blvd, Oakland CA 94612. CleanPro will provide commercial "
            "office cleaning including daily janitorial and weekly deep cleaning for $3,500 per month, paid monthly. "
            "Starting February 1, 2026. 30 days notice to terminate. California law, mediation for disputes. Client owns all IP."
        ),
        "expected_extracted": {
            "service_type": "Standard",
            "client_name": "GlobalTech Solutions",
            "client_address": "500 Market St, San Francisco CA 94105",
            "contractor_name": "CleanPro Services LLC",
            "contractor_address": "200 Oak Blvd, Oakland CA 94612",
            "services_description": "commercial office cleaning including daily janitorial and weekly deep cleaning",
            "effective_date": "February 1, 2026",
            "compensation_amount": "$3,500 per month",
            "payment_schedule": "Monthly",
            "termination_notice_days": "30",
            "governing_law": "California",
            "ip_ownership": "Client",
            "dispute_resolution_method": "Mediation",
        },
        "expected_follow_ups": [],
        "acceptable_extra_extracted": ["service_type"],
    },
    {
        "id": 2,
        "name": "Partial service agreement - missing addresses",
        "contract_type": "ServiceAgreement",
        "prompt": (
            "I need a service agreement. WebWorks Design will build a new website for Summit Corp. "
            "The project costs $15,000 total, to be paid upon completion. Use Texas law."
        ),
        "expected_extracted": {
            "service_type": "Standard",
            "client_name": "Summit Corp",
            "contractor_name": "WebWorks Design",
            "services_description": "build a new website",
            "compensation_amount": "$15,000",
            "payment_schedule": "Upon Completion",
            "governing_law": "Texas",
        },
        "expected_follow_ups": [
            "client_address",
            "contractor_address",
            "effective_date",
            "termination_notice_days",
            "ip_ownership",
        ],
        "acceptable_extra_extracted": ["service_type"],
    },
    {
        "id": 3,
        "name": "Conversational service agreement",
        "contract_type": "ServiceAgreement",
        "prompt": (
            "We're hiring a landscaping company to maintain our office grounds. "
            "Green Thumb Landscaping will do weekly lawn care and seasonal planting for $800 a month. "
            "We're in Florida."
        ),
        "expected_extracted": {
            "contractor_name": "Green Thumb Landscaping",
            "services_description": "weekly lawn care and seasonal planting",
            "compensation_amount": "$800",
            "governing_law": "Florida",
        },
        "expected_follow_ups": [
            "client_name",
            "client_address",
            "contractor_address",
            "effective_date",
            "termination_notice_days",
            "ip_ownership",
        ],
        "acceptable_extra_extracted": ["service_type", "payment_schedule"],
    },
    {
        "id": 4,
        "name": "Near-empty service agreement prompt",
        "contract_type": "ServiceAgreement",
        "prompt": "I need a service contract for some IT consulting work.",
        "expected_extracted": {},
        "expected_follow_ups": [
            "client_name",
            "client_address",
            "contractor_name",
            "contractor_address",
            "services_description",
            "effective_date",
            "compensation_amount",
            "payment_schedule",
            "termination_notice_days",
            "governing_law",
            "ip_ownership",
        ],
        "acceptable_extra_extracted": ["service_type", "services_description"],
        "must_not_extract": [
            "client_name",
            "client_address",
            "contractor_name",
            "contractor_address",
            "effective_date",
            "compensation_amount",
            "termination_notice_days",
            "governing_law",
        ],
    },
]


# ---------------------------------------------------------------------------
# Value comparison
# ---------------------------------------------------------------------------
def normalize_for_comparison(value) -> str:
    if value is None:
        return ""
    s = str(value).strip().lower()
    s = " ".join(s.split())
    s = s.rstrip(".")
    s = s.replace("the state of ", "")
    if s.endswith(" law"):
        s = s[:-4].strip()
    return s


def values_match(expected, actual) -> bool:
    e = normalize_for_comparison(expected)
    a = normalize_for_comparison(actual)
    if not e or not a:
        return e == a
    if e == a:
        return True
    if e in a or a in e:
        return True
    try:
        if float(e) == float(a):
            return True
    except (ValueError, TypeError):
        pass
    return False


# ---------------------------------------------------------------------------
# Follow-up question quality check
# ---------------------------------------------------------------------------
def check_question_quality(field_name_str: str, question: str) -> dict:
    """Check if a follow-up question is clear and answerable."""
    issues = []

    if not question or len(question.strip()) < 5:
        issues.append("question is empty or too short")

    if question and not question.strip().endswith("?") and not question.strip().endswith(")"):
        issues.append("question doesn't end with ? or option list")

    # Check that question references something relevant to the field
    field_keywords = field_name_str.replace("_", " ").lower().split()
    question_lower = question.lower()
    has_relevance = any(kw in question_lower for kw in field_keywords if len(kw) > 2)
    if not has_relevance:
        issues.append(f"question may not relate to field '{field_name_str}'")

    return {
        "field": field_name_str,
        "question": question,
        "quality": "good" if not issues else "poor",
        "issues": issues,
    }


# ---------------------------------------------------------------------------
# Scoring
# ---------------------------------------------------------------------------
def score_test(test: dict, verified_answers: dict, follow_ups: List[dict]) -> dict:
    expected = test["expected_extracted"]
    expected_fups = set(test.get("expected_follow_ups", []))
    must_not_extract = set(test.get("must_not_extract", []))
    acceptable_extra = set(test.get("acceptable_extra_extracted", []))
    actual_fup_fields = {item["field"] for item in follow_ups}

    results = {
        "test_id": test["id"],
        "test_name": test["name"],
        "field_results": {},
        "follow_up_results": {},
        "follow_up_quality": [],
    }

    # --- Extraction scoring ---
    total_expected = len(expected)
    correct_fields = 0
    field_details = {}

    for fname, expected_val in expected.items():
        actual_val = verified_answers.get(fname)
        match = values_match(expected_val, actual_val)
        if match:
            correct_fields += 1
        field_details[fname] = {
            "expected": expected_val,
            "actual": actual_val,
            "match": match,
        }

    # --- Hallucination detection ---
    hallucinated = []
    for fname in verified_answers:
        if fname in must_not_extract:
            hallucinated.append(fname)

    # --- Follow-up coverage (the critical metric) ---
    fup_total = len(expected_fups)
    fup_correct = len(expected_fups & actual_fup_fields)
    fup_missing = expected_fups - actual_fup_fields

    fup_details = {}
    for fname in expected_fups:
        fup_details[fname] = {
            "expected": "must be asked",
            "was_asked": fname in actual_fup_fields,
        }

    # --- Follow-up question quality ---
    fup_quality = []
    for item in follow_ups:
        quality = check_question_quality(item["field"], item["question"])
        fup_quality.append(quality)

    good_questions = sum(1 for q in fup_quality if q["quality"] == "good")
    total_questions = len(fup_quality) if fup_quality else 1

    # --- Scoring weights ---
    # Extraction: 30% (conservative is OK)
    # Follow-up coverage: 50% (critical — must catch all missing fields)
    # No hallucination: 20% (must not fabricate)
    extraction_score = correct_fields / total_expected if total_expected > 0 else 1.0
    followup_coverage = fup_correct / fup_total if fup_total > 0 else 1.0
    hallucination_score = 1.0 - (len(hallucinated) * 0.25)  # 25% penalty each
    hallucination_score = max(0, hallucination_score)
    question_quality_score = good_questions / total_questions

    overall = (
        extraction_score * 0.30
        + followup_coverage * 0.50
        + hallucination_score * 0.20
    )

    results["field_results"] = field_details
    results["follow_up_results"] = fup_details
    results["follow_up_quality"] = fup_quality
    results["hallucinated_fields"] = hallucinated
    results["follow_up_missing"] = list(fup_missing)
    results["extraction_score"] = round(extraction_score, 3)
    results["followup_coverage"] = round(followup_coverage, 3)
    results["hallucination_score"] = round(hallucination_score, 3)
    results["question_quality_score"] = round(question_quality_score, 3)
    results["overall_score"] = round(overall, 3)

    return results


# ---------------------------------------------------------------------------
# Assembly test (end-to-end)
# ---------------------------------------------------------------------------
NDA_FOLLOW_UP_ANSWERS = {
    "nda_type": "Unilateral",
    "party_a_name": "Acme Corp.",
    "party_a_entity_details": "Delaware corporation",
    "party_a_email": "legal@acme.com",
    "party_b_name": "Beta Ventures LLC",
    "party_b_entity_details": "Connecticut LLC",
    "party_b_email": "ops@betaventures.com",
    "purpose": "evaluating a potential partnership",
    "confidentiality_period_number": 3,
    "confidentiality_period_unit": "years",
    "governing_law": "Connecticut",
    "dispute_resolution_method": "Litigation",
}

CONSULTING_FOLLOW_UP_ANSWERS = {
    "consulting_type": "Standard",
    "client_name": "Acme Corp",
    "client_address": "123 Main St, Hartford CT 06103",
    "client_email": "legal@acme.com",
    "consultant_name": "Jane Smith Consulting",
    "consultant_address": "456 Oak Ave, New Haven CT 06510",
    "consultant_email": "jane@smithconsulting.com",
    "effective_date": "January 15, 2026",
    "services_description": "marketing strategy and brand positioning services",
    "compensation_amount": "$25,000",
    "payment_schedule": "Net 30",
    "term_end_date": "June 30, 2026",
    "termination_notice_days": 30,
    "governing_law": "Connecticut",
    "ip_ownership": "Client",
    "dispute_resolution_method": "Arbitration",
}

EMPLOYMENT_FOLLOW_UP_ANSWERS = {
    "employment_type": "Standard",
    "employer_name": "TechCorp Inc.",
    "employer_address": "100 Innovation Dr, San Francisco, CA 94105",
    "employee_name": "Jane Smith",
    "employee_address": "456 Oak Ave, San Francisco, CA 94110",
    "job_title": "Senior Software Engineer",
    "job_duties": "Design, develop, and maintain software applications",
    "start_date": "January 15, 2026",
    "employment_basis": "Full-Time",
    "compensation_amount": "$120,000 per year",
    "pay_frequency": "Bi-Weekly",
    "work_location": "100 Innovation Dr, San Francisco, CA 94105",
    "termination_notice_days": 30,
    "governing_law": "California",
    "dispute_resolution_method": "Arbitration",
}

SERVICE_FOLLOW_UP_ANSWERS = {
    "service_type": "Standard",
    "client_name": "GlobalTech Solutions Inc.",
    "client_address": "500 Market St, San Francisco CA 94105",
    "contractor_name": "CleanPro Services LLC",
    "contractor_address": "200 Oak Blvd, Oakland CA 94612",
    "services_description": "commercial office cleaning services",
    "effective_date": "February 1, 2026",
    "compensation_amount": "$3,500 per month",
    "payment_schedule": "Monthly",
    "termination_notice_days": 30,
    "governing_law": "California",
    "ip_ownership": "Client",
    "dispute_resolution_method": "Mediation",
}

FOLLOW_UP_ANSWERS_BY_TYPE = {
    "NDA": NDA_FOLLOW_UP_ANSWERS,
    "ConsultingAgreement": CONSULTING_FOLLOW_UP_ANSWERS,
    "EmploymentAgreement": EMPLOYMENT_FOLLOW_UP_ANSWERS,
    "ServiceAgreement": SERVICE_FOLLOW_UP_ANSWERS,
}


def fill_missing_and_assemble(verified_answers: dict, follow_ups: List[dict], contract_type: str = "NDA") -> tuple:
    fallback = FOLLOW_UP_ANSWERS_BY_TYPE.get(contract_type, NDA_FOLLOW_UP_ANSWERS)
    for item in follow_ups:
        fname = item["field"]
        if fname not in verified_answers and fname in fallback:
            verified_answers[fname] = fallback[fname]

    verified_answers = add_derived_defaults(verified_answers, contract_type)
    try:
        contract_text, unresolved, _ = assemble_contract(verified_answers, contract_type)
        return contract_text, unresolved
    except Exception as e:
        return f"ASSEMBLY ERROR: {e}", ["ERROR"]


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def run_tests(test_ids: Optional[List[int]] = None, contract_type: str = "NDA"):
    if contract_type == "all":
        # Run all contract type tests
        print("\n" + "=" * 60)
        print("  RUNNING NDA TESTS")
        print("=" * 60)
        nda_results, nda_avg = run_tests(test_ids=test_ids, contract_type="NDA")

        print("\n" + "=" * 60)
        print("  RUNNING CONSULTING TESTS")
        print("=" * 60)
        ca_results, ca_avg = run_tests(test_ids=test_ids, contract_type="ConsultingAgreement")

        print("\n" + "=" * 60)
        print("  RUNNING EMPLOYMENT TESTS")
        print("=" * 60)
        ea_results, ea_avg = run_tests(test_ids=test_ids, contract_type="EmploymentAgreement")

        print("\n" + "=" * 60)
        print("  RUNNING SERVICE TESTS")
        print("=" * 60)
        sa_results, sa_avg = run_tests(test_ids=test_ids, contract_type="ServiceAgreement")

        print(f"\n{'='*60}")
        print("COMBINED SUMMARY")
        print(f"{'='*60}")
        print(f"  NDA average:         {nda_avg*100:.1f}%")
        print(f"  Consulting average:  {ca_avg*100:.1f}%")
        print(f"  Employment average:  {ea_avg*100:.1f}%")
        print(f"  Service average:     {sa_avg*100:.1f}%")
        combined = (nda_avg + ca_avg + ea_avg + sa_avg) / 4
        print(f"  Combined average:    {combined*100:.1f}%")
        return nda_results + ca_results + ea_results + sa_results, combined

    test_bank_map = {
        "NDA": TESTS,
        "ConsultingAgreement": CONSULTING_TESTS,
        "EmploymentAgreement": EMPLOYMENT_TESTS,
        "ServiceAgreement": SERVICE_TESTS,
    }
    test_bank = test_bank_map.get(contract_type, TESTS)
    tests_to_run = test_bank if test_ids is None else [t for t in test_bank if t["id"] in test_ids]
    ct = contract_type  # shorthand

    all_results = []
    total_score = 0

    for test in tests_to_run:
        test_ct = test.get("contract_type", ct)
        print(f"\n{'='*60}")
        print(f"[{test_ct}] TEST {test['id']}: {test['name']}")
        print(f"{'='*60}")
        print(f"Prompt: {test['prompt'][:100]}...")

        start = time.time()
        print("  Extracting...")
        extraction = extract_answers_from_prompt(test["prompt"], test_ct)
        extract_time = time.time() - start

        raw_answers = extraction.get("known_answers", {})
        raw_answers = {k: v for k, v in raw_answers.items() if v is not None and str(v).strip() != ""}
        print(f"  Extracted {len(raw_answers)} fields in {extract_time:.1f}s")

        verify_start = time.time()
        verified_answers, follow_ups, evidence = verify_and_prepare(extraction, test_ct)
        verify_time = time.time() - verify_start
        total_time = time.time() - start

        print(f"  Verified {len(verified_answers)} fields, {len(follow_ups)} follow-ups")

        # Score
        result = score_test(test, verified_answers, follow_ups)
        result["contract_type"] = test_ct
        result["extract_time"] = round(extract_time, 1)
        result["total_time"] = round(total_time, 1)

        # Assembly
        contract_text, unresolved = fill_missing_and_assemble(dict(verified_answers), list(follow_ups), test_ct)
        result["assembly_success"] = len(unresolved) == 0
        result["unresolved_placeholders"] = unresolved

        all_results.append(result)
        total_score += result["overall_score"]

        # Print results
        print(f"\n  Verified answers: {json.dumps(verified_answers, indent=4, default=str)}")

        print(f"\n  Follow-ups ({len(follow_ups)}):")
        for fup in follow_ups:
            quality = next((q for q in result["follow_up_quality"] if q["field"] == fup["field"]), {})
            qmark = "OK" if quality.get("quality") == "good" else "!!"
            print(f"    [{qmark}] {fup['field']}: \"{fup['question']}\"")
            if quality.get("issues"):
                for issue in quality["issues"]:
                    print(f"         ^ {issue}")

        if result["hallucinated_fields"]:
            print(f"\n  !! HALLUCINATED: {result['hallucinated_fields']}")
        if result["follow_up_missing"]:
            print(f"  !! MISSING FOLLOW-UPS: {result['follow_up_missing']}")

        print(f"\n  Extraction:     {result['extraction_score']*100:.0f}% (weight 30%)")
        print(f"  Follow-up cov:  {result['followup_coverage']*100:.0f}% (weight 50%)")
        print(f"  No hallucinate: {result['hallucination_score']*100:.0f}% (weight 20%)")
        print(f"  Question qual:  {result['question_quality_score']*100:.0f}%")
        print(f"  Overall score:  {result['overall_score']*100:.0f}%")
        print(f"  Assembly:       {'PASS' if result['assembly_success'] else 'FAIL'}")
        if unresolved:
            print(f"  Unresolved:     {unresolved}")
        print(f"  Time:           {total_time:.1f}s")

        # Field misses
        for fname, detail in result["field_results"].items():
            if not detail["match"]:
                print(f"  MISS: {fname} expected={detail['expected']!r} got={detail['actual']!r}")

    # Summary
    avg_score = total_score / len(tests_to_run) if tests_to_run else 0
    assembly_pass = sum(1 for r in all_results if r["assembly_success"])

    # Count follow-up coverage across all tests
    total_expected_fups = 0
    total_actual_fups = 0
    total_hallucinations = 0
    for r in all_results:
        for fup_detail in r["follow_up_results"].values():
            total_expected_fups += 1
            if fup_detail["was_asked"]:
                total_actual_fups += 1
        total_hallucinations += len(r["hallucinated_fields"])

    print(f"\n{'='*60}")
    print(f"SUMMARY ({contract_type})")
    print(f"{'='*60}")
    for r in all_results:
        status = "PASS" if r["overall_score"] >= 0.8 else "FAIL"
        asm = "ASM:OK" if r["assembly_success"] else "ASM:FAIL"
        fup_status = f"FUP:{r['followup_coverage']*100:.0f}%"
        print(f"  Test {r['test_id']}: {r['overall_score']*100:5.1f}%  {status}  {asm}  {fup_status}  ({r['total_time']:.0f}s)  {r['test_name']}")

    print(f"\n  Average score:       {avg_score*100:.1f}%")
    print(f"  Assembly pass:       {assembly_pass}/{len(tests_to_run)}")
    print(f"  Follow-up coverage:  {total_actual_fups}/{total_expected_fups} ({total_actual_fups/total_expected_fups*100:.0f}%)" if total_expected_fups > 0 else "  Follow-up coverage:  N/A")
    print(f"  Hallucinations:      {total_hallucinations}")
    print(f"  Target:              95%+")

    # Save results
    suffix = f"_{contract_type.lower()}" if contract_type != "NDA" else ""
    output_path = ROOT / "output" / f"test_results{suffix}.json"
    output_path.write_text(json.dumps(all_results, indent=2, default=str), encoding="utf-8")
    print(f"\n  Results saved to {output_path}")

    return all_results, avg_score


def get_test_bank(contract_type: str) -> List[dict]:
    return {
        "NDA": TESTS,
        "ConsultingAgreement": CONSULTING_TESTS,
        "EmploymentAgreement": EMPLOYMENT_TESTS,
        "ServiceAgreement": SERVICE_TESTS,
    }[contract_type]


def resolve_contract_types(contract_type: str) -> List[str]:
    return ALL_CONTRACT_TYPES if contract_type == "all" else [contract_type]


def complete_answers_for_type(contract_type: str) -> Dict[str, Any]:
    return add_derived_defaults(dict(FOLLOW_UP_ANSWERS_BY_TYPE[contract_type]), contract_type)


def merge_fallback_answers(
    verified_answers: Dict[str, Any],
    follow_ups: List[dict],
    contract_type: str,
    evidence: Optional[Dict[str, str]] = None,
) -> tuple[Dict[str, Any], Dict[str, str]]:
    answers = dict(verified_answers)
    final_evidence = dict(evidence or {})
    fallback = FOLLOW_UP_ANSWERS_BY_TYPE[contract_type]
    for item in follow_ups:
        fname = item["field"]
        if fname not in answers and fname in fallback:
            answers[fname] = fallback[fname]
            final_evidence.setdefault(
                fname, f"Test fixture follow-up answer for {fname}"
            )
    return add_derived_defaults(answers, contract_type), final_evidence


def save_named_results(name: str, payload: dict) -> Path:
    path = OUTPUT_DIR / f"{name}.json"
    path.write_text(json.dumps(payload, indent=2, default=str), encoding="utf-8")
    return path


def execute_case(name: str, func, contract_type: Optional[str] = None) -> dict:
    started = time.time()
    try:
        details = func() or {}
        passed = True
        error = None
    except Exception as exc:
        details = {}
        passed = False
        error = f"{type(exc).__name__}: {exc}"

    duration = round(time.time() - started, 2)
    return {
        "name": name,
        "contract_type": contract_type,
        "passed": passed,
        "duration_seconds": duration,
        "details": details,
        "error": error,
    }


def print_custom_suite_summary(title: str, results: List[dict]) -> None:
    passed = sum(1 for item in results if item["passed"])
    failed = len(results) - passed
    print(f"\n{'=' * 60}")
    print(f"{title}")
    print(f"{'=' * 60}")
    for item in results:
        status = "PASS" if item["passed"] else "FAIL"
        scope = f"[{item['contract_type']}] " if item.get("contract_type") else ""
        print(f"  {status:4}  {scope}{item['name']}  ({item['duration_seconds']:.2f}s)")
        if item["error"]:
            print(f"        {item['error']}")
    print(f"\n  Passed: {passed}/{len(results)}")
    print(f"  Failed: {failed}")


def validate_artifacts(
    docx_path: Path,
    sidecar_path: Path,
    contract_text: str,
    clauses: List[dict],
    answers: Dict[str, Any],
    contract_type: str,
) -> dict:
    doc = DocxDocument(str(docx_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    sidecar = json.loads(sidecar_path.read_text(encoding="utf-8"))

    assert docx_path.exists() and docx_path.stat().st_size > 0, "DOCX was not created"
    assert sidecar_path.exists() and sidecar_path.stat().st_size > 0, "Sidecar was not created"
    assert len(paragraphs) >= 10, f"Expected at least 10 non-empty paragraphs, got {len(paragraphs)}"
    assert "{{" not in contract_text and "}}" not in contract_text, "Unresolved placeholders remained in contract text"
    assert "Title" in styles, "DOCX does not contain a Title paragraph"
    assert any(style == "Heading 1" for style in styles), "DOCX does not contain a Heading 1 paragraph"
    assert isinstance(sidecar.get("clauses"), list) and sidecar["clauses"], "Citation sidecar is missing clause entries"
    assert len(sidecar["clauses"]) == len(clauses), "Citation sidecar clause count mismatch"

    full_text = "\n".join(paragraphs)
    for field_name in PARTY_FIELD_MAP.get(contract_type, []):
        party_name = str(answers.get(field_name, "")).strip()
        if party_name:
            assert party_name in full_text, f"Expected '{party_name}' in DOCX body"

    return {
        "docx_path": str(docx_path),
        "sidecar_path": str(sidecar_path),
        "paragraph_count": len(paragraphs),
        "clause_count": len(clauses),
    }


def clause_candidate_counts(
    contract_type: str,
    subtype: str,
    subtype_field: str,
) -> Dict[str, int]:
    library, _, _, _ = load_resources(contract_type)
    counts: Dict[str, int] = {}
    for clause in library:
        if clause.get(subtype_field) == subtype:
            clause_name = clause["clause_name"]
            counts[clause_name] = counts.get(clause_name, 0) + 1
    return counts


def assert_rag_metadata_shape(
    metadata: Dict[str, dict],
    ordered_clause_names: List[str],
    candidate_counts: Dict[str, int],
) -> dict:
    assert len(metadata) == len(ordered_clause_names), "Missing metadata for one or more ordered clauses"

    rag_clauses = []
    deterministic_clauses = []
    for clause_name in ordered_clause_names:
        entry = metadata[clause_name]
        num_candidates = candidate_counts.get(clause_name, 0)
        if num_candidates <= 1:
            assert entry["method"] == "deterministic", f"{clause_name} should be deterministic"
            assert entry["score"] == 1.0, f"{clause_name} should have a deterministic score of 1.0"
            deterministic_clauses.append(clause_name)
        else:
            assert entry["method"] == "rag", f"{clause_name} should use RAG"
            assert 0.0 <= entry["score"] <= 1.0, f"{clause_name} returned an invalid similarity score"
            rag_clauses.append(clause_name)

    return {
        "rag_clause_count": len(rag_clauses),
        "deterministic_clause_count": len(deterministic_clauses),
    }


def run_end_to_end_suite(contract_type: str = "all") -> dict:
    results = []
    for current_type in resolve_contract_types(contract_type):
        primary_test = get_test_bank(current_type)[0]

        def _case() -> dict:
            extraction = extract_answers_from_prompt(primary_test["prompt"], current_type)
            verified_answers, follow_ups, evidence = verify_and_prepare(extraction, current_type)
            final_answers, final_evidence = merge_fallback_answers(
                verified_answers, follow_ups, current_type, evidence
            )

            assembled = assemble_contract(final_answers, current_type, use_rag=True)
            contract_text, unresolved, clauses, rag_meta = (
                assembled if len(assembled) == 4 else (*assembled, None)
            )

            assert not unresolved, f"Unresolved placeholders: {unresolved}"

            docx_path, sidecar_path = generate_contract_docx(
                contract_text,
                clauses,
                rag_meta,
                final_answers,
                final_evidence,
                current_type,
                current_type,
            )
            artifact_details = validate_artifacts(
                docx_path, sidecar_path, contract_text, clauses, final_answers, current_type
            )
            artifact_details["initial_follow_up_count"] = len(follow_ups)
            return artifact_details

        results.append(execute_case("end_to_end_pipeline", _case, current_type))

    payload = {
        "suite": "e2e",
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "results": results,
        "passed": sum(1 for item in results if item["passed"]),
        "failed": sum(1 for item in results if not item["passed"]),
    }
    save_named_results("test_results_e2e", payload)
    print_custom_suite_summary("END-TO-END SUITE", results)
    return payload


def run_rag_suite(contract_type: str = "all") -> dict:
    results = []

    def _run_case(current_type: str, answers_a: Dict[str, Any], answers_b: Optional[Dict[str, Any]] = None) -> dict:
        final_a = add_derived_defaults(dict(answers_a), current_type)
        library, order, _, config = load_resources(current_type)
        subtype_field = config.get("subtype_field", "nda_type")
        subtype = final_a[subtype_field]
        ordered_clause_names = order[subtype]
        candidate_counts = clause_candidate_counts(current_type, subtype, subtype_field)

        selected_a, metadata_a = select_clauses_rag(
            current_type, subtype, subtype_field, ordered_clause_names, final_a
        )
        summary = assert_rag_metadata_shape(metadata_a, ordered_clause_names, candidate_counts)
        assert len(selected_a) == len(ordered_clause_names), "Not every ordered clause was selected"
        summary["selected_clause_count"] = len(selected_a)

        if answers_b is not None:
            final_b = add_derived_defaults(dict(answers_b), current_type)
            selected_b, metadata_b = select_clauses_rag(
                current_type, subtype, subtype_field, ordered_clause_names, final_b
            )
            assert len(selected_b) == len(ordered_clause_names), "Context-B selection missed clauses"
            changed = [
                clause_name
                for clause_name in ordered_clause_names
                if metadata_a[clause_name]["variant_id"] != metadata_b[clause_name]["variant_id"]
            ]
            score_changed = [
                clause_name
                for clause_name in ordered_clause_names
                if metadata_a[clause_name]["score"] != metadata_b[clause_name]["score"]
            ]
            summary["variant_changes"] = changed
            summary["variant_change_count"] = len(changed)
            summary["score_change_count"] = len(score_changed)
            summary["context_sensitive_variant_selection"] = bool(changed)

        return summary

    selected_types = resolve_contract_types(contract_type)
    for current_type in selected_types:
        if current_type == "NDA":
            answers = complete_answers_for_type("NDA")
            results.append(
                execute_case(
                    "rag_metadata_shape_and_determinism",
                    lambda: _run_case("NDA", answers),
                    "NDA",
                )
            )
            continue

        answers_a = complete_answers_for_type(current_type)
        answers_b = dict(answers_a)
        if current_type == "ConsultingAgreement":
            answers_b["dispute_resolution_method"] = "Litigation"
            answers_b["ip_ownership"] = "Consultant"
            case_name = "rag_context_sensitivity"
        elif current_type == "EmploymentAgreement":
            answers_b["dispute_resolution_method"] = "Litigation"
            answers_b["employment_basis"] = "Part-Time"
            case_name = "rag_context_sensitivity"
        else:
            answers_b["dispute_resolution_method"] = "Litigation"
            answers_b["ip_ownership"] = "Contractor"
            case_name = "rag_context_sensitivity"

        results.append(
            execute_case(
                case_name,
                lambda ct=current_type, a=answers_a, b=answers_b: _run_case(ct, a, b),
                current_type,
            )
        )

    payload = {
        "suite": "rag",
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "results": results,
        "passed": sum(1 for item in results if item["passed"]),
        "failed": sum(1 for item in results if not item["passed"]),
    }
    save_named_results("test_results_rag", payload)
    print_custom_suite_summary("RAG SUITE", results)
    return payload


def run_follow_up_suite() -> dict:
    pending = [
        {"field": "nda_type", "question": "Should this be a Mutual or Unilateral NDA?"},
        {"field": "party_a_entity_details", "question": "What are the entity details for Party A?"},
        {"field": "party_b_entity_details", "question": "What are the entity details for Party B?"},
    ]
    results = []

    def _assert_follow_up_parse(answer_text: str, expected_fields: List[str], expected_remaining: List[str]) -> dict:
        parsed = parse_follow_up_answers(answer_text, pending, "NDA")
        missing = [item["field"] for item in pending if item["field"] not in parsed]
        assert sorted(parsed.keys()) == sorted(expected_fields), "Unexpected parsed follow-up fields"
        assert sorted(missing) == sorted(expected_remaining), "Unexpected remaining follow-up fields"
        return {"parsed_fields": sorted(parsed.keys()), "remaining_fields": missing}

    results.append(
        execute_case(
            "numbered_answers_complete",
            lambda: _assert_follow_up_parse(
                "1. Unilateral\n2. Delaware corporation\n3. Connecticut LLC",
                ["nda_type", "party_a_entity_details", "party_b_entity_details"],
                [],
            ),
        )
    )
    results.append(
        execute_case(
            "labeled_answers_complete",
            lambda: _assert_follow_up_parse(
                "NDA Type: Mutual\nParty A Entity Details: Delaware corporation\nParty B Entity Details: Connecticut LLC",
                ["nda_type", "party_a_entity_details", "party_b_entity_details"],
                [],
            ),
        )
    )
    results.append(
        execute_case(
            "line_by_line_answers_complete",
            lambda: _assert_follow_up_parse(
                "Mutual\nDelaware corporation\nConnecticut LLC",
                ["nda_type", "party_a_entity_details", "party_b_entity_details"],
                [],
            ),
        )
    )
    results.append(
        execute_case(
            "partial_answers_leave_remaining_fields",
            lambda: _assert_follow_up_parse(
                "1. Unilateral\n2. Delaware corporation",
                ["nda_type", "party_a_entity_details"],
                ["party_b_entity_details"],
            ),
        )
    )

    payload = {
        "suite": "followup",
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "results": results,
        "passed": sum(1 for item in results if item["passed"]),
        "failed": sum(1 for item in results if not item["passed"]),
    }
    save_named_results("test_results_followup", payload)
    print_custom_suite_summary("FOLLOW-UP SUITE", results)
    return payload


def run_docx_suite(contract_type: str = "all") -> dict:
    results = []

    for current_type in resolve_contract_types(contract_type):
        def _case() -> dict:
            final_answers = complete_answers_for_type(current_type)
            evidence = {
                key: f'Test fixture value for "{key}"'
                for key, value in final_answers.items()
                if value is not None and str(value).strip() != ""
            }
            assembled = assemble_contract(final_answers, current_type, use_rag=True)
            contract_text, unresolved, clauses, rag_meta = (
                assembled if len(assembled) == 4 else (*assembled, None)
            )
            assert not unresolved, f"Unresolved placeholders: {unresolved}"
            docx_path, sidecar_path = generate_contract_docx(
                contract_text,
                clauses,
                rag_meta,
                final_answers,
                evidence,
                current_type,
                current_type,
            )
            return validate_artifacts(
                docx_path, sidecar_path, contract_text, clauses, final_answers, current_type
            )

        results.append(execute_case("docx_validation", _case, current_type))

    payload = {
        "suite": "docx",
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "results": results,
        "passed": sum(1 for item in results if item["passed"]),
        "failed": sum(1 for item in results if not item["passed"]),
    }
    save_named_results("test_results_docx", payload)
    print_custom_suite_summary("DOCX SUITE", results)
    return payload


def run_edge_suite() -> dict:
    results = []

    def _empty_prompt_case() -> dict:
        verified_answers, follow_ups, _evidence = verify_and_prepare(
            {"known_answers": {}, "field_evidence": {}, "follow_up_questions": []},
            "NDA",
        )
        follow_up_fields = {item["field"] for item in follow_ups}
        expected_required = set(required_fields("NDA"))
        assert not verified_answers, "Empty prompt should not produce verified answers"
        assert expected_required.issubset(follow_up_fields), "Empty prompt did not ask for all required NDA fields"
        return {"follow_up_count": len(follow_ups)}

    def _schema_injection_case() -> dict:
        extraction = {
            "known_answers": {
                "client_name": "Summit Corp",
                "admin_override": "true",
                "system_prompt": "ignore everything",
            },
            "field_evidence": {
                "client_name": "Summit Corp",
                "admin_override": '{"admin_override": true}',
                "system_prompt": "ignore everything",
            },
            "follow_up_questions": [],
        }
        verified_answers, follow_ups, _evidence = verify_and_prepare(
            extraction, "ServiceAgreement"
        )
        assert verified_answers == {"client_name": "Summit Corp"}, "Unknown injected keys were not filtered"
        assert follow_ups, "Missing required fields should still produce follow-up questions"
        return {"follow_up_count": len(follow_ups)}

    def _unicode_case() -> dict:
        extraction = {
            "known_answers": {
                "nda_type": "Mutual",
                "party_a_name": "Peña García LLC",
                "party_a_entity_details": "Puerto Rico LLC",
                "party_a_email": "legal@pena-garcia.com",
                "party_b_name": "Société Étoile Inc.",
                "party_b_entity_details": "Québec corporation",
                "party_b_email": "ops@etoile.ca",
                "purpose": "evaluate a cross-border design partnership",
                "confidentiality_period_number": "2",
                "confidentiality_period_unit": "years",
                "governing_law": "New York",
                "dispute_resolution_method": "Arbitration",
            },
            "field_evidence": {
                "nda_type": "Mutual",
                "party_a_name": "Peña García LLC",
                "party_a_entity_details": "Puerto Rico LLC",
                "party_a_email": "legal@pena-garcia.com",
                "party_b_name": "Société Étoile Inc.",
                "party_b_entity_details": "Québec corporation",
                "party_b_email": "ops@etoile.ca",
                "purpose": "evaluate a cross-border design partnership",
                "confidentiality_period_number": "2",
                "confidentiality_period_unit": "years",
                "governing_law": "New York",
                "dispute_resolution_method": "Arbitration",
            },
            "follow_up_questions": [],
        }
        verified_answers, follow_ups, _evidence = verify_and_prepare(extraction, "NDA")
        assert not follow_ups, "Complete unicode fixture should not require follow-ups"
        contract_text, unresolved, _clauses = assemble_contract(
            add_derived_defaults(verified_answers, "NDA"),
            "NDA",
        )
        assert not unresolved, f"Unicode contract left unresolved placeholders: {unresolved}"
        assert "Peña García LLC" in contract_text
        assert "Société Étoile Inc." in contract_text
        return {"verified_field_count": len(verified_answers)}

    def _long_prompt_case() -> dict:
        long_services = "enterprise website redesign and accessibility review " * 45
        extraction = {
            "known_answers": {
                "service_type": "Standard",
                "client_name": "Summit Corp",
                "client_address": "100 King St, Austin TX 78701",
                "contractor_name": "WebWorks Design",
                "contractor_address": "200 Elm St, Dallas TX 75201",
                "services_description": long_services.strip(),
                "effective_date": "May 1, 2026",
                "compensation_amount": "$15,000",
                "payment_schedule": "Upon Completion",
                "termination_notice_days": "15",
                "governing_law": "Texas",
                "ip_ownership": "Client",
                "dispute_resolution_method": "Litigation",
            },
            "field_evidence": {
                "service_type": "Standard",
                "client_name": "Summit Corp",
                "client_address": "100 King St, Austin TX 78701",
                "contractor_name": "WebWorks Design",
                "contractor_address": "200 Elm St, Dallas TX 75201",
                "services_description": long_services.strip(),
                "effective_date": "May 1, 2026",
                "compensation_amount": "$15,000",
                "payment_schedule": "Upon Completion",
                "termination_notice_days": "15",
                "governing_law": "Texas",
                "ip_ownership": "Client",
                "dispute_resolution_method": "Litigation",
            },
            "follow_up_questions": [],
        }
        verified_answers, follow_ups, _evidence = verify_and_prepare(extraction, "ServiceAgreement")
        contract_text, unresolved, _clauses = assemble_contract(
            add_derived_defaults(verified_answers, "ServiceAgreement"),
            "ServiceAgreement",
        )
        assert len(long_services) > 2000, "Fixture text is not actually long enough"
        assert not unresolved, f"Long-text fixture left unresolved placeholders: {unresolved}"
        assert long_services.strip()[:80] in contract_text, "Long services text was not preserved"
        return {
            "prompt_length": len(long_services),
            "verified_field_count": len(verified_answers),
            "follow_up_count": len(follow_ups),
        }

    def _missing_evidence_case() -> dict:
        extraction = {
            "known_answers": {
                "nda_type": "Mutual",
                "party_a_name": "Acme Corp",
                "party_b_name": "Beta Ventures LLC",
            },
            "field_evidence": {
                "nda_type": "Mutual",
                "party_a_name": "",
                "party_b_name": "Beta Ventures LLC",
            },
            "follow_up_questions": [],
        }
        verified_answers, follow_ups, _evidence = verify_and_prepare(extraction, "NDA")
        follow_up_fields = {item["field"] for item in follow_ups}
        assert verified_answers.get("party_b_name"), "Contradictory prompt should still retain Party B"
        assert "party_a_name" not in verified_answers, "Answer without evidence should not be verified"
        assert "party_a_name" in follow_up_fields, "Missing-evidence field should come back as a follow-up"
        return {
            "verified_field_count": len(verified_answers),
            "follow_up_count": len(follow_ups),
        }

    def _zero_value_case() -> dict:
        extraction = {
            "known_answers": {
                "service_type": "Standard",
                "client_name": "Zero Budget Labs",
                "client_address": "1 Test Way, Boston MA 02110",
                "contractor_name": "Starter Studio LLC",
                "contractor_address": "2 Sample Rd, Boston MA 02111",
                "services_description": "prototype review and onboarding support",
                "effective_date": "April 1, 2026",
                "compensation_amount": "$0",
                "payment_schedule": "Upon Completion",
                "termination_notice_days": "0",
                "governing_law": "Massachusetts",
                "ip_ownership": "Client",
                "dispute_resolution_method": "Litigation",
            },
            "field_evidence": {
                "service_type": "Standard",
                "client_name": "Zero Budget Labs",
                "client_address": "1 Test Way, Boston MA 02110",
                "contractor_name": "Starter Studio LLC",
                "contractor_address": "2 Sample Rd, Boston MA 02111",
                "services_description": "prototype review and onboarding support",
                "effective_date": "April 1, 2026",
                "compensation_amount": "$0",
                "payment_schedule": "Upon Completion",
                "termination_notice_days": "0",
                "governing_law": "Massachusetts",
                "ip_ownership": "Client",
                "dispute_resolution_method": "Litigation",
            },
            "follow_up_questions": [],
        }
        verified_answers, follow_ups, _evidence = verify_and_prepare(
            extraction, "ServiceAgreement"
        )
        assert not follow_ups, "Complete zero-value fixture should not require follow-ups"
        assert str(verified_answers["compensation_amount"]) == "$0", "Zero-dollar compensation was lost"
        assert str(verified_answers["termination_notice_days"]) == "0", "Zero-day notice was lost"
        return {
            "compensation_amount": verified_answers["compensation_amount"],
            "termination_notice_days": verified_answers["termination_notice_days"],
        }

    results.extend(
        [
            execute_case("empty_prompt_requires_followups", _empty_prompt_case, "NDA"),
            execute_case("schema_injection_unknown_keys_filtered", _schema_injection_case, "ServiceAgreement"),
            execute_case("unicode_values_preserved", _unicode_case, "NDA"),
            execute_case("very_long_text_preserved", _long_prompt_case, "ServiceAgreement"),
            execute_case("missing_evidence_values_rejected", _missing_evidence_case, "NDA"),
            execute_case("zero_value_numerics_preserved", _zero_value_case, "ServiceAgreement"),
        ]
    )

    payload = {
        "suite": "edge",
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "results": results,
        "passed": sum(1 for item in results if item["passed"]),
        "failed": sum(1 for item in results if not item["passed"]),
    }
    save_named_results("test_results_edge", payload)
    print_custom_suite_summary("EDGE-CASE SUITE", results)
    return payload


def run_full_suite() -> dict:
    extraction_results, extraction_average = run_tests(contract_type="all")
    e2e_payload = run_end_to_end_suite("all")
    rag_payload = run_rag_suite("all")
    edge_payload = run_edge_suite()
    followup_payload = run_follow_up_suite()
    docx_payload = run_docx_suite("all")

    summary = {
        "suite": "full",
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "extraction_average": round(extraction_average, 3),
        "extraction_tests": len(extraction_results),
        "e2e_passed": e2e_payload["passed"],
        "e2e_failed": e2e_payload["failed"],
        "rag_passed": rag_payload["passed"],
        "rag_failed": rag_payload["failed"],
        "edge_passed": edge_payload["passed"],
        "edge_failed": edge_payload["failed"],
        "followup_passed": followup_payload["passed"],
        "followup_failed": followup_payload["failed"],
        "docx_passed": docx_payload["passed"],
        "docx_failed": docx_payload["failed"],
    }
    save_named_results(
        "test_results_full",
        {
            "summary": summary,
            "extraction_results": extraction_results,
            "e2e": e2e_payload,
            "rag": rag_payload,
            "edge": edge_payload,
            "followup": followup_payload,
            "docx": docx_payload,
        },
    )
    print(f"\n{'=' * 60}")
    print("FULL SUITE SUMMARY")
    print(f"{'=' * 60}")
    print(f"  Extraction average: {extraction_average * 100:.1f}%")
    print(f"  E2E:                {e2e_payload['passed']}/{len(e2e_payload['results'])}")
    print(f"  RAG:                {rag_payload['passed']}/{len(rag_payload['results'])}")
    print(f"  Edge:               {edge_payload['passed']}/{len(edge_payload['results'])}")
    print(f"  Follow-up:          {followup_payload['passed']}/{len(followup_payload['results'])}")
    print(f"  DOCX:               {docx_payload['passed']}/{len(docx_payload['results'])}")
    return summary


if __name__ == "__main__":
    contract_type = "NDA"
    test_ids = None
    mode = "extraction"

    for arg in sys.argv[1:]:
        lower = arg.lower()
        if arg.isdigit():
            if test_ids is None:
                test_ids = []
            test_ids.append(int(arg))
        elif lower in ("consulting", "consultingagreement"):
            contract_type = "ConsultingAgreement"
        elif lower in ("employment", "employmentagreement"):
            contract_type = "EmploymentAgreement"
        elif lower in ("service", "serviceagreement"):
            contract_type = "ServiceAgreement"
        elif lower == "nda":
            contract_type = "NDA"
        elif lower == "all":
            contract_type = "all"
        elif lower in {"e2e", "rag", "edge", "followup", "docx", "full"}:
            mode = lower

    if mode == "e2e":
        run_end_to_end_suite(contract_type)
    elif mode == "rag":
        run_rag_suite(contract_type)
    elif mode == "edge":
        run_edge_suite()
    elif mode == "followup":
        run_follow_up_suite()
    elif mode == "docx":
        run_docx_suite(contract_type)
    elif mode == "full":
        run_full_suite()
    else:
        run_tests(test_ids=test_ids, contract_type=contract_type)
