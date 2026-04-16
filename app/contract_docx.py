"""
DOCX contract generator.

Produces a clean, professional Word document from the already-assembled
``contract_text`` string (the same text the CLI and Web UI see). Nothing
about citations, sources, or highlights lives in the body of the draft —
the DOCX reads like a normal legal document.

Citation/audit data is written to a companion JSON sidecar next to the
DOCX so an interactive viewer (see README: Interactive Citations) can
render it separately.
"""

import json
import re
import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# Document styling
# ---------------------------------------------------------------------------
def setup_styles(doc: Document):
    """Configure document styles for a professional legal document."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    title_style = doc.styles["Title"]
    title_style.font.name = "Times New Roman"
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(18)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(11)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


# ---------------------------------------------------------------------------
# Plain-text → DOCX renderer
# ---------------------------------------------------------------------------
_HEADING_NUMBERED_RE = re.compile(r"^\d+\.\s+[A-Z][A-Z\s/&\-]+$")
_SUBHEADING_PREFIX_RE = re.compile(r"^(?:[A-Z]\.\s+|\([A-Za-z0-9]\)\s+)")


def _is_all_caps_heading(line: str) -> bool:
    """An ALL CAPS short line that isn't a placeholder or signature rule."""
    if len(line) < 3 or len(line) > 80:
        return False
    if line.startswith("{{") or line.startswith("_") or line.startswith("$"):
        return False
    letters = [c for c in line if c.isalpha()]
    if not letters:
        return False
    return all(c.isupper() for c in letters)


def _is_subheading(line: str) -> bool:
    """``A. Definitions`` is a sub-heading; ``A. The Client is of the opinion...``
    is a recital (normal paragraph). Heuristic: must start with the prefix,
    be short (<= 60 chars), and not look like a full sentence.
    """
    if not _SUBHEADING_PREFIX_RE.match(line):
        return False
    if len(line) > 60:
        return False
    body = _SUBHEADING_PREFIX_RE.sub("", line).strip()
    if not body:
        return False
    # Full sentences end with a period and contain spaces → paragraph, not heading
    if body.endswith(".") and " " in body:
        return False
    return True


def render_contract_text_to_docx(doc: Document, contract_text: str):
    """Parse the assembled plain-text contract and emit DOCX paragraphs.

    One source line = one DOCX paragraph. Blank lines are skipped (Word
    handles spacing via paragraph styles). Classification:
      - First ALL CAPS line → document Title
      - Subsequent ALL CAPS line or ``1. HEADING`` → Heading 1
      - Short label-like ``A. Definitions`` / ``(a) Scope`` → Heading 2
      - Everything else → body paragraph
    """
    title_emitted = False

    for raw in contract_text.split("\n"):
        stripped = raw.strip()
        if not stripped:
            continue

        if _HEADING_NUMBERED_RE.match(stripped):
            doc.add_heading(stripped, level=1)
            continue

        if _is_all_caps_heading(stripped):
            if not title_emitted:
                doc.add_paragraph(stripped, style="Title")
                title_emitted = True
            else:
                doc.add_heading(stripped, level=1)
            continue

        if _is_subheading(stripped):
            doc.add_heading(stripped, level=2)
            continue

        doc.add_paragraph(stripped)


# ---------------------------------------------------------------------------
# Citation sidecar (for interactive viewer)
# ---------------------------------------------------------------------------
def build_citation_sidecar(
    clauses: List[dict],
    rag_metadata: Optional[Dict[str, dict]],
    verified_answers: Dict[str, Any],
    evidence: Dict[str, str],
    contract_type: str,
    label: str,
) -> dict:
    """Package clause sources + user-input evidence into a JSON-ready dict."""
    clause_entries = []
    for i, clause in enumerate(clauses, 1):
        cname = clause.get("clause_name", "?")
        meta = (rag_metadata or {}).get(cname, {})
        clause_entries.append(
            {
                "index": i,
                "clause_name": cname,
                "source": clause.get("source", "Unknown"),
                "variant_id": clause.get("variant_id", "?"),
                "method": meta.get("method", "default"),
                "num_candidates": meta.get("num_candidates", 1),
                "score": meta.get("score", None),
                "text": clause.get("text", ""),
            }
        )

    evidence_entries = []
    for field_name, value in verified_answers.items():
        if not value or field_name == "special_provisions":
            continue
        evidence_entries.append(
            {
                "field": field_name,
                "label": field_name.replace("_", " ").title(),
                "value": str(value),
                "evidence": str(evidence.get(field_name, "User-provided (follow-up)")),
            }
        )

    return {
        "contract_type": contract_type,
        "label": label,
        "generated_at": datetime.datetime.now().isoformat(timespec="seconds"),
        "clauses": clause_entries,
        "user_input_evidence": evidence_entries,
    }


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------
def generate_contract_docx(
    contract_text: str,
    clauses: List[dict],
    rag_metadata: Optional[Dict[str, dict]],
    verified_answers: Dict[str, Any],
    evidence: Dict[str, str],
    contract_type: str,
    label: str,
) -> Tuple[Path, Path]:
    """Generate a clean DOCX + citation JSON sidecar.

    Returns (docx_path, sidecar_json_path).
    """
    doc = Document()
    setup_styles(doc)
    render_contract_text_to_docx(doc, contract_text)

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_type = contract_type.lower().replace(" ", "_")
    base = f"{safe_type}_{timestamp}"

    docx_path = OUTPUT_DIR / f"{base}.docx"
    doc.save(str(docx_path))

    sidecar = build_citation_sidecar(
        clauses, rag_metadata, verified_answers, evidence, contract_type, label
    )
    sidecar_path = OUTPUT_DIR / f"{base}.citations.json"
    sidecar_path.write_text(json.dumps(sidecar, indent=2))

    return docx_path, sidecar_path


# ---------------------------------------------------------------------------
# Quick test
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import sys
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from assemble_contract import assemble_contract

    answers = {
        "service_type": "Standard",
        "client_name": "GlobalTech Solutions Inc.",
        "client_address": "500 Market St, San Francisco, CA 94105",
        "contractor_name": "CleanPro Services LLC",
        "contractor_address": "200 Oak Blvd, Oakland, CA 94612",
        "services_description": "Commercial office cleaning including daily janitorial, weekly deep cleaning, and monthly floor maintenance.",
        "effective_date": "February 1, 2026",
        "compensation_amount": "$3,500 per month",
        "payment_schedule": "Monthly",
        "termination_notice_days": "30",
        "governing_law": "California",
        "ip_ownership": "Client",
        "special_provisions": "",
    }

    result = assemble_contract(answers, "ServiceAgreement", use_rag=True)
    contract_text, unresolved, clauses, rag_meta = (
        result if len(result) == 4 else (*result, None)
    )
    evidence = {k: f'User said: "{v}"' for k, v in answers.items() if v}

    docx_path, sidecar_path = generate_contract_docx(
        contract_text, clauses, rag_meta, answers, evidence,
        "ServiceAgreement", "Service Agreement",
    )
    print(f"DOCX:     {docx_path}")
    print(f"Sidecar:  {sidecar_path}")
